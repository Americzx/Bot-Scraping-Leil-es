import os
import requests
from bs4 import BeautifulSoup
from docx import Document

# Coloque o URL do anúncio do megaleiloes aqui:
url = 'https://www.megaleiloes.com.br/imoveis/terrenos-e-lotes/sp/itapolis/terreno-253-m2-jardim-estoril-itapolis-sp-j96140'
response = requests.get(url)

if response.status_code != 200:
    print('Falha ao conectar')
    exit()

soup = BeautifulSoup(response.text, 'html.parser')

document = Document()
document.add_heading('Detalhes do Anúncio', level=1)

data = {}

title = soup.find('h1', class_='section-header')
data['Título'] = title.text.strip() if title else 'N/A'

localiza = soup.find('div', class_='locality item')
data['Localização'] = localiza.text.strip() if localiza else 'N/A'

processo = soup.find('div', class_='process-number item')
data['Número do processo'] = processo.text.strip() if processo else 'N/A'

data['Link'] = url 

firstactive = soup.find('div', class_='instance first active')
if firstactive:
    data['Primeiro leilão'] = firstactive.text.strip() 
else:
    firstinactive = soup.find('div', class_='instance first passed')
    data['Primeiro leilão inativo'] = firstinactive.text.strip()

secondactive = soup.find('div', class_='instance active')
if secondactive:
    data['Segundo leilão'] = secondactive.text.strip() 
else:
    secondinactive = soup.find('div', class_='instance passed')
    data['Segundo leilão inativo'] = secondinactive.text.strip()  

auctioneer = soup.find('div', class_='author item')
data['Leiloeiro'] = auctioneer.text.strip() if auctioneer else 'N/A'

description = soup.find('div', class_='description')
data['Descrição'] = description.text.strip() if description else 'N/A'

div_ft = soup.find_all('div', class_='page')
links_imagens = []

for div in div_ft:
    imagens = div.find_all('img')
    links_imagens.extend([img['src'] for img in imagens])

data['Imagens'] = links_imagens if links_imagens else 'N/A'

# Especifica o caminho para a pasta onde o documento será salvo
pasta_especifica = r'C:\Users\kayke\Documents\resultados'

# Verifica se a pasta existe, se não, a cria
if not os.path.exists(pasta_especifica):
    os.makedirs(pasta_especifica)

# Adiciona os dados ao documento
for key, value in data.items():
    document.add_paragraph(f'{key}: {value}')

# Salva o documento na pasta específica
caminho_documento = os.path.join(pasta_especifica, 'casa_anuncio.docx')
document.save(caminho_documento)
print(f'Documento salvo em: {caminho_documento}')

print('kayke')
